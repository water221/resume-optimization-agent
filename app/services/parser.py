import io
import re
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader


def normalize_text_line(line: str) -> str:
    return re.sub(r"\s+", "", line or "")


def dedupe_lines(lines: list[str]) -> list[str]:
    deduped: list[str] = []
    seen_keys: list[str] = []
    for line in lines:
        raw = (line or "").strip()
        if not raw:
            continue
        key = normalize_text_line(raw)
        if not key:
            continue

        is_duplicate = False
        for sk in seen_keys:
            if key == sk:
                is_duplicate = True
                break
            if len(key) >= 20 and len(sk) >= 20 and (key in sk or sk in key):
                is_duplicate = True
                break
        if is_duplicate:
            continue

        seen_keys.append(key)
        deduped.append(raw)

    filtered: list[str] = []
    keys = [normalize_text_line(x) for x in deduped]
    for idx, line in enumerate(deduped):
        key = keys[idx]
        if len(key) > 80:
            contained_count = 0
            for j, other_key in enumerate(keys):
                if j == idx:
                    continue
                if 6 <= len(other_key) <= 40 and other_key in key:
                    contained_count += 1
                if contained_count >= 3:
                    break
            if contained_count >= 3:
                continue
        filtered.append(line)

    return filtered


def extract_docx_text_from_xml(raw: bytes) -> str:
    lines: list[str] = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            candidates = [
                name for name in zf.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]
            for name in candidates:
                xml_bytes = zf.read(name)
                root = ET.fromstring(xml_bytes)
                for p in root.findall(".//w:p", ns):
                    texts = [t.text for t in p.findall(".//w:t", ns) if t.text]
                    line = "".join(texts).strip()
                    if line:
                        lines.append(line)
    except Exception:
        return ""

    return "\n".join(dedupe_lines(lines))


def read_upload(file: UploadFile) -> str:
    filename = (file.filename or "").lower()
    raw = file.file.read()
    if not raw:
        return ""

    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(raw))
        paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]

        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            tables_text.append(p.text.strip())

        combined = paragraphs_text + tables_text

        xml_text = extract_docx_text_from_xml(raw)
        if xml_text:
            existing_keys = {normalize_text_line(x) for x in combined if x.strip()}
            for line in xml_text.splitlines():
                line = line.strip()
                key = normalize_text_line(line)
                if line and key and key not in existing_keys:
                    combined.append(line)
                    existing_keys.add(key)

        return "\n".join(dedupe_lines(combined))

    if filename.endswith(".md") or filename.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")

    return raw.decode("utf-8", errors="ignore")
