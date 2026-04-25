import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


def sanitize_base_filename(name: str) -> str:
    base = Path((name or "").strip()).stem
    if not base:
        base = "optimized_resume"
    base = re.sub(r"[\\/:*?\"<>|]+", "_", base)
    base = re.sub(r"\s+", "_", base).strip("._")
    return base or "optimized_resume"


def build_download_filename(session: Dict[str, Any], fmt: str) -> str:
    source_name = session.get("source_filename") or "optimized_resume"
    base = sanitize_base_filename(source_name)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{base}_{ts}.{fmt}"


def clean_markdown_inline(text: str) -> str:
    text = strip_problematic_chars(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text


def parse_markdown_line(line: str) -> tuple[str, str]:
    stripped = line.strip()
    if not stripped:
        return "empty", ""
    if re.fullmatch(r"(?:-{3,}|\*{3,}|_{3,}|—{3,}|–{3,}|(?:-\s*){3,})", stripped):
        return "hr", ""

    heading_match = re.match(r"^(#{1,6})\s*(.+)$", stripped)
    if heading_match:
        level = len(heading_match.group(1))
        return f"h{level}", heading_match.group(2).strip()

    bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
    if bullet_match:
        body = bullet_match.group(1).strip()
        nested_heading = re.match(r"^(#{1,6})\s*(.+)$", body)
        if nested_heading:
            level = len(nested_heading.group(1))
            return f"h{level}", nested_heading.group(2).strip()
        return "bullet", body

    numbered_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
    if numbered_match:
        body = numbered_match.group(1).strip()
        nested_heading = re.match(r"^(#{1,6})\s*(.+)$", body)
        if nested_heading:
            level = len(nested_heading.group(1))
            return f"h{level}", nested_heading.group(2).strip()
        return "numbered", body
    return "text", line


def register_pdf_font() -> str:
    candidate_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]
    for i, font_path in enumerate(candidate_paths):
        if Path(font_path).exists():
            font_name = f"CJKFont{i}"
            pdfmetrics.registerFont(TTFont(font_name, font_path))
            return font_name
    return "Helvetica"


def build_pdf_font_map(base_font_name: str) -> Dict[str, str]:
    return {"regular": base_font_name, "bold": base_font_name}


def strip_problematic_chars(text: str) -> str:
    if not text:
        return ""
    cleaned_chars: list[str] = []
    for ch in text:
        if ch in {"\u200b", "\ufeff", "↩", "\ufffd"}:
            continue
        cat = unicodedata.category(ch)
        if cat in {"Cc", "Cf", "Cs", "Co", "Cn"}:
            continue
        cleaned_chars.append(ch)
    return "".join(cleaned_chars)


def remove_square_noise(text: str) -> str:
    if not text:
        return ""
    cleaned = re.sub(r"[■□▪▫▮▯█▓▒▉▊▋▌▍▎▏]{3,}", "", text)
    cleaned = re.sub(r"[\u25A0-\u25FF]{4,}", "", cleaned)
    return cleaned


def normalize_export_lines(content: str) -> list[str]:
    raw_lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    normalized: list[str] = []
    prev_empty = True
    for raw in raw_lines:
        line = strip_problematic_chars(raw)
        line = remove_square_noise(line)
        line = line.replace("\t", " ").rstrip()

        dense_square_count = len(re.findall(r"[■□▪▫▮▯█▓▒]", line))
        compact = re.sub(r"\s+", "", line)
        if compact and dense_square_count >= 6 and dense_square_count / max(1, len(compact)) > 0.25:
            continue

        if not line.strip():
            if not prev_empty and normalized:
                normalized.append("")
            prev_empty = True
            continue
        normalized.append(line)
        prev_empty = False

    while normalized and not normalized[-1].strip():
        normalized.pop()
    return normalized


def sanitize_pdf_text(text: str) -> str:
    text = clean_markdown_inline(text)
    text = remove_square_noise(text)
    text = text.replace("•", "·")
    text = re.sub(r"[\U00010000-\U0010FFFF]", "", text)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    filtered: list[str] = []
    for ch in text:
        cp = ord(ch)
        if 0x25A0 <= cp <= 0x25FF or cp in {0x2B1B, 0x2B1C}:
            continue
        if unicodedata.category(ch) == "So":
            continue
        filtered.append(ch)
    return "".join(filtered)


def strip_unsupported_pdf_glyphs(text: str, font_name: str) -> str:
    if not text:
        return ""
    try:
        font = pdfmetrics.getFont(font_name)
        face = getattr(font, "face", None)
        char_to_glyph = getattr(face, "charToGlyph", None)
        if not isinstance(char_to_glyph, dict):
            return text
    except Exception:
        return text

    out_chars: list[str] = []
    for ch in text:
        cp = ord(ch)
        if ch in {" ", "\t"}:
            out_chars.append(ch)
            continue
        if cp in char_to_glyph:
            out_chars.append(ch)
    return "".join(out_chars)


def apply_docx_default_styles(doc: Document):
    def set_style(style_name: str, size: int, bold: bool = False):
        try:
            style = doc.styles[style_name]
            style.font.name = "Microsoft YaHei"
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = RGBColor(0, 0, 0)
            try:
                rpr = style.element.get_or_add_rPr()
                rfonts = rpr.get_or_add_rFonts()
                rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            except Exception:
                pass
        except Exception:
            pass

    def set_heading_style(style_name: str, size: int):
        try:
            style = doc.styles[style_name]
            style.font.name = "Microsoft YaHei"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(47, 84, 150)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(2)
            try:
                rpr = style.element.get_or_add_rPr()
                rfonts = rpr.get_or_add_rFonts()
                rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            except Exception:
                pass
        except Exception:
            pass

    set_style("Normal", 11, False)
    set_style("List Bullet", 11, False)
    set_heading_style("Heading 1", 16)
    set_heading_style("Heading 2", 14)
    set_heading_style("Heading 3", 13)
    set_heading_style("Heading 4", 12)


def apply_docx_run_style(run, *, bold: bool = False, color: tuple[int, int, int] = (0, 0, 0)):
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*color)
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    except Exception:
        pass


def wrap_text_for_pdf(text: str, font_name: str, font_size: int, max_width: float) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for ch in text:
        test = current + ch
        if pdfmetrics.stringWidth(test, font_name, font_size) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def render_markdown_to_docx(doc: Document, content: str):
    apply_docx_default_styles(doc)
    for line in normalize_export_lines(content):
        line_type, body = parse_markdown_line(line)
        if line_type in {"empty", "hr"}:
            continue
        if line_type.startswith("h") and line_type[1:].isdigit():
            heading_level = min(int(line_type[1:]), 4)
            doc.add_heading(clean_markdown_inline(body), level=heading_level)
            continue
        if line_type in {"bullet", "numbered"}:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(clean_markdown_inline(body))
            apply_docx_run_style(r, bold=False, color=(0, 0, 0))
            continue

        p = doc.add_paragraph()
        r = p.add_run(clean_markdown_inline(line))
        apply_docx_run_style(r, bold=False, color=(0, 0, 0))


def render_markdown_to_pdf(c: canvas.Canvas, content: str, width: float, height: float, font_map: Dict[str, str]):
    x, y = 40, height - 45
    right_margin = 40
    max_width = width - x - right_margin
    line_height = 16

    def ensure_page(next_step: int = line_height):
        nonlocal y
        if y < 45 + next_step:
            c.showPage()
            y = height - 45

    for line in normalize_export_lines(content):
        line_type, body = parse_markdown_line(line)
        if line_type == "empty":
            y -= 10
            continue

        if line_type == "hr":
            ensure_page(12)
            c.setStrokeColor(HexColor("#9CA3AF"))
            c.line(x, y, width - right_margin, y)
            y -= 14
            continue

        if line_type == "h1":
            ensure_page(20)
            c.setFont(font_map["bold"], 14)
            text = strip_unsupported_pdf_glyphs(sanitize_pdf_text(body), font_map["bold"])
            if not text.strip():
                continue
            for chunk in wrap_text_for_pdf(text, font_map["bold"], 14, max_width):
                ensure_page(20)
                c.drawString(x, y, chunk)
                y -= 20
            continue

        if line_type in {"h2", "h3", "h4", "h5", "h6"}:
            ensure_page(18)
            c.setFont(font_map["bold"], 12)
            text = strip_unsupported_pdf_glyphs(sanitize_pdf_text(body), font_map["bold"])
            if not text.strip():
                continue
            for chunk in wrap_text_for_pdf(text, font_map["bold"], 12, max_width):
                ensure_page(18)
                c.drawString(x, y, chunk)
                y -= 18
            continue

        text = sanitize_pdf_text(body if line_type in {"bullet", "numbered"} else line)
        if line_type == "bullet":
            text = f"· {text}"
        if line_type == "numbered":
            text = f"1) {text}"
        text = strip_unsupported_pdf_glyphs(text, font_map["regular"])
        if not text.strip():
            continue

        c.setFont(font_map["regular"], 10)
        for chunk in wrap_text_for_pdf(text, font_map["regular"], 10, max_width):
            ensure_page(line_height)
            c.drawString(x, y, chunk)
            y -= line_height


def safe_filename(session_id: str, ext: str, export_dir: Path) -> Path:
    return export_dir / f"optimized_resume_{session_id}.{ext}"


def export_md(content: str, session: Dict[str, Any], session_id: str, export_dir: Path):
    path = safe_filename(session_id, "md", export_dir)
    path.write_text(content, encoding="utf-8")
    return path, build_download_filename(session, "md"), "text/markdown"


def export_docx(content: str, session: Dict[str, Any], session_id: str, export_dir: Path):
    path = safe_filename(session_id, "docx", export_dir)
    doc = Document()
    render_markdown_to_docx(doc, content)
    doc.save(path)
    return path, build_download_filename(session, "docx"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def export_pdf(content: str, session: Dict[str, Any], session_id: str, export_dir: Path):
    path = safe_filename(session_id, "pdf", export_dir)
    font_name = register_pdf_font()
    font_map = build_pdf_font_map(font_name)
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    render_markdown_to_pdf(c, content, width, height, font_map)
    c.save()
    return path, build_download_filename(session, "pdf"), "application/pdf"
