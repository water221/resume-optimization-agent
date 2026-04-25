import io
import json
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, PlainTextResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
from pydantic import BaseModel, Field
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

BASE_DIR = Path(__file__).resolve().parent
EXPORT_DIR = BASE_DIR / "exports"
EXPORT_DIR.mkdir(exist_ok=True)

app = FastAPI(title="Resume Optimization Agent", version="1.0.0")
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")

# Simple server-side memory. In production, replace it with Redis/Postgres.
SESSIONS: Dict[str, Dict[str, Any]] = {}


class AnalyzeRequest(BaseModel):
    resume_text: str = Field(..., min_length=20)
    jd_text: str = Field(..., min_length=20)


class OptimizeRequest(BaseModel):
    session_id: str
    user_confirmed: bool = True
    extra_instruction: Optional[str] = ""


def get_client() -> OpenAI:
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="Missing DEEPSEEK_API_KEY in environment variables.")
    return OpenAI(api_key=api_key, base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com"))


def llm_chat(messages: list[dict[str, str]], temperature: float = 0.2) -> str:
    client = get_client()
    model = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
    resp = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
    )
    return resp.choices[0].message.content or ""


def trim_text(text: str, max_chars: int = 30000) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[内容过长，已截断。请上传更精简版本以获得更完整分析。]"


def read_upload(file: UploadFile) -> str:
    filename = (file.filename or "").lower()
    raw = file.file.read()
    if not raw:
        return ""

    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if filename.endswith(".md") or filename.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")

    # Best effort for unknown plain text-like files.
    return raw.decode("utf-8", errors="ignore")


def extract_json(text: str) -> Dict[str, Any]:
    cleaned = text.strip()
    cleaned = re.sub(r"^```json\s*", "", cleaned)
    cleaned = re.sub(r"^```\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    match = re.search(r"\{.*\}", cleaned, flags=re.S)
    if not match:
        raise ValueError("LLM did not return JSON.")
    return json.loads(match.group(0))


ANALYZE_SYSTEM = """你是一个严谨的简历优化 Agent，目标是帮助求职者基于真实经历提升岗位匹配度。\n\n硬性规则：\n1. 只能基于原始简历中已经出现的事实进行分析，不允许编造学历、公司、项目、指标、技术栈。\n2. 如果 JD 需要但简历没有证据，必须放入主要缺口，而不是替用户虚构。\n3. 输出必须是合法 JSON，不要输出 Markdown 代码块。\n4. 建议要具体到简历修改动作，例如“把 X 项目改写为面向 Y 场景的 Agent 项目描述”。\n"""

OPTIMIZE_SYSTEM = """你是一个严谨的中文简历改写 Agent。\n\n硬性规则：\n1. 必须忠于原始简历事实，不得新增不存在的经历、公司、奖项、论文、指标。\n2. 可以重组表达、强化与 JD 相关的关键词、突出已有项目中的职责/技术/结果。\n3. 对不确定信息不要写成确定事实。\n4. 输出一份可直接复制到 Markdown/Word 的中文简历正文。\n5. 使用清晰标题、项目经历 bullet、专业技能分组，适合 AI Agent/AI 应用开发岗位。\n"""


def build_analysis_prompt(resume_text: str, jd_text: str) -> str:
    return f"""
请分析下面的【原始简历】与【目标 JD】的匹配情况。\n\n请输出 JSON，字段如下：\n{{\n  "match_score": 0-100的整数,\n  "target_role_summary": "用一句话概括JD核心要求",\n  "matching_highlights": ["匹配亮点1", "匹配亮点2"],\n  "main_gaps": ["主要缺口1", "主要缺口2"],\n  "optimization_suggestions": [\n    {{"area": "模块名称", "problem": "当前问题", "action": "具体修改动作", "evidence_from_resume": "来自原简历的事实依据"}}\n  ],\n  "risk_warnings": ["哪些内容不能乱写或需要用户补充确认"]\n}}\n\n【目标 JD】\n{jd_text}\n\n【原始简历】\n{resume_text}\n"""


def build_optimize_prompt(session: Dict[str, Any], extra_instruction: str = "") -> str:
    return f"""
用户已经确认根据分析结果优化简历。请基于以下内容生成优化后的简历。\n\n【目标 JD】\n{session['jd_text']}\n\n【原始简历】\n{session['resume_text']}\n\n【上一轮分析结果】\n{json.dumps(session['analysis'], ensure_ascii=False, indent=2)}\n\n【用户补充要求】\n{extra_instruction or '无'}\n\n请输出优化后的完整简历正文。\n"""


@app.get("/", response_class=HTMLResponse)
def index() -> str:
    return (BASE_DIR / "static" / "index.html").read_text(encoding="utf-8")


@app.post("/api/parse-resume")
def parse_resume(file: UploadFile = File(...)):
    text = trim_text(read_upload(file))
    if not text:
        raise HTTPException(status_code=400, detail="无法从文件中解析出文本，请尝试粘贴简历文本。")
    return {"resume_text": text}


@app.post("/api/analyze")
def analyze(req: AnalyzeRequest):
    resume_text = trim_text(req.resume_text)
    jd_text = trim_text(req.jd_text, max_chars=20000)
    raw = llm_chat([
        {"role": "system", "content": ANALYZE_SYSTEM},
        {"role": "user", "content": build_analysis_prompt(resume_text, jd_text)},
    ])
    try:
        analysis = extract_json(raw)
    except Exception:
        analysis = {
            "match_score": None,
            "target_role_summary": "模型未返回标准 JSON，以下为原始分析结果。",
            "matching_highlights": [],
            "main_gaps": [],
            "optimization_suggestions": [],
            "risk_warnings": ["请检查模型输出格式。"],
            "raw_output": raw,
        }
    session_id = str(uuid.uuid4())
    SESSIONS[session_id] = {
        "resume_text": resume_text,
        "jd_text": jd_text,
        "analysis": analysis,
        "optimized_resume": None,
        "created_at": datetime.utcnow().isoformat(),
    }
    return {"session_id": session_id, "analysis": analysis}


@app.post("/api/optimize")
def optimize(req: OptimizeRequest):
    if not req.user_confirmed:
        raise HTTPException(status_code=400, detail="用户尚未确认生成优化简历。")
    session = SESSIONS.get(req.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found. Please analyze again.")
    optimized = llm_chat([
        {"role": "system", "content": OPTIMIZE_SYSTEM},
        {"role": "user", "content": build_optimize_prompt(session, req.extra_instruction or "")},
    ])
    session["optimized_resume"] = optimized.strip()
    return {"session_id": req.session_id, "optimized_resume": session["optimized_resume"]}


def safe_filename(session_id: str, ext: str) -> Path:
    return EXPORT_DIR / f"optimized_resume_{session_id}.{ext}"


def get_resume_or_404(session_id: str) -> str:
    session = SESSIONS.get(session_id)
    if not session or not session.get("optimized_resume"):
        raise HTTPException(status_code=404, detail="Optimized resume not found. Please generate it first.")
    return session["optimized_resume"]


@app.get("/api/export/{session_id}.{fmt}")
def export_resume(session_id: str, fmt: str):
    content = get_resume_or_404(session_id)
    fmt = fmt.lower()
    if fmt == "md":
        path = safe_filename(session_id, "md")
        path.write_text(content, encoding="utf-8")
        return FileResponse(path, filename=path.name, media_type="text/markdown")

    if fmt == "docx":
        path = safe_filename(session_id, "docx")
        doc = Document()
        for line in content.splitlines():
            if line.startswith("# "):
                doc.add_heading(line.replace("# ", ""), level=1)
            elif line.startswith("## "):
                doc.add_heading(line.replace("## ", ""), level=2)
            elif line.strip().startswith("- "):
                doc.add_paragraph(line.strip()[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(path)
        return FileResponse(path, filename=path.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if fmt == "pdf":
        path = safe_filename(session_id, "pdf")
        font_path = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
        font_name = "NotoSansCJK"
        if Path(font_path).exists():
            pdfmetrics.registerFont(TTFont(font_name, font_path))
        else:
            font_name = "Helvetica"
        c = canvas.Canvas(str(path), pagesize=A4)
        width, height = A4
        x, y = 40, height - 45
        c.setFont(font_name, 10)
        for paragraph in content.splitlines():
            if not paragraph.strip():
                y -= 10
                continue
            # naive wrapping, sufficient for basic validation
            chunks = [paragraph[i:i + 42] for i in range(0, len(paragraph), 42)]
            for chunk in chunks:
                if y < 45:
                    c.showPage()
                    c.setFont(font_name, 10)
                    y = height - 45
                c.drawString(x, y, chunk)
                y -= 16
        c.save()
        return FileResponse(path, filename=path.name, media_type="application/pdf")

    raise HTTPException(status_code=400, detail="Unsupported format. Use md, docx, or pdf.")


@app.get("/health")
def health():
    return {"status": "ok"}


@app.exception_handler(Exception)
def global_exception_handler(_, exc: Exception):
    if isinstance(exc, HTTPException):
        raise exc
    return JSONResponse(status_code=500, content={"detail": str(exc)})
